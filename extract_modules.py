#!/usr/bin/env python3
"""
Script to extract content from .docx module files and prepare it for integration
into the RiddleNet learning platform.
"""

import os
import re
from docx import Document
from pathlib import Path

def extract_text_from_docx(file_path):
    """Extract all text content from a .docx file."""
    try:
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    content.append(" | ".join(row_data))
        
        return "\n".join(content)
    except Exception as e:
        print(f"Error reading {file_path}: {str(e)}")
        return None

def process_module_content(text):
    """Process extracted text into structured content suitable for web display."""
    if not text:
        return None
    
    # Split into sections based on common patterns
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Identify headings (common patterns)
        if (line.isupper() and len(line) > 5) or \
           (line.startswith('Module ') and ':' in line) or \
           (line.startswith('Lesson ') and ':' in line) or \
           (line.startswith('Chapter ') and ':' in line):
            processed_lines.append(f"<h3>{line}</h3>")
        elif line.endswith(':') and len(line.split()) <= 5:
            processed_lines.append(f"<h4>{line}</h4>")
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            # List items
            clean_line = line[1:].strip()
            processed_lines.append(f"<li>{clean_line}</li>")
        elif re.match(r'^\d+\.', line):
            # Numbered list items
            processed_lines.append(f"<li>{line}</li>")
        else:
            # Regular paragraph
            processed_lines.append(f"<p>{line}</p>")
    
    return "\n".join(processed_lines)

def convert_to_html_lesson(title, content):
    """Convert processed content into the HTML format used by the learning platform."""
    html_content = f'''
                <div class="lesson-content">
                    <h2>{title}</h2>
                    
                    <div class="lesson-section">
                        {content}
                    </div>
                    
                    <div class="info-box">
                        <h4>Key Concept</h4>
                        <p>This content has been extracted from the module documentation. Interactive elements and quizzes will be added in future updates.</p>
                    </div>
                </div>
            '''
    return html_content

def main():
    """Main function to process all module files."""
    modules_dir = Path("modules")
    output_file = "extracted_module_content.py"
    
    if not modules_dir.exists():
        print(f"Modules directory not found: {modules_dir}")
        return
    
    # Find all .docx files
    docx_files = list(modules_dir.glob("*.docx"))
    if not docx_files:
        print("No .docx files found in modules directory")
        return
    
    print(f"Found {len(docx_files)} module files:")
    for file in docx_files:
        print(f"  - {file.name}")
    
    # Process each file
    module_data = {}
    
    for file_path in sorted(docx_files):
        print(f"\nProcessing {file_path.name}...")
        
        # Extract module number/id from filename
        filename = file_path.stem  # Remove .docx extension
        
        # Try to extract module identifier
        if "Module-1" in filename and "ITEP-207" in filename:
            module_id = "1.1"
            title = "Introduction to Computer Networks"
        elif "Module-2.1" in filename:
            module_id = "2.1"
            title = "Data Link Layer Fundamentals"
        elif "Module-2" in filename and "Module-2.1" not in filename:
            module_id = "2.2"
            title = "Data Link Layer Protocols"
        elif "Module-3" in filename:
            module_id = "3.1"
            title = "Network Layer Fundamentals"
        elif "Module-4.1" in filename:
            module_id = "4.1"
            title = "Transport Layer - TCP/UDP"
        elif "Module-4.2" in filename:
            module_id = "4.2"
            title = "Application Layer Protocols"
        elif "Module-4.3" in filename:
            module_id = "4.3"
            title = "Network Applications and Services"
        elif "Module-4" in filename and "Module-4." not in filename:
            module_id = "4.4"
            title = "Advanced Network Applications"
        else:
            # Fallback - use filename
            module_id = filename.replace("-", ".").replace("Module.", "").replace("ITEP.207.Networking.1", "").strip(".")
            title = f"Module {module_id}"
        
        # Extract content
        raw_content = extract_text_from_docx(file_path)
        if raw_content:
            processed_content = process_module_content(raw_content)
            if processed_content:
                html_content = convert_to_html_lesson(title, processed_content)
                module_data[module_id] = {
                    "title": title,
                    "content": html_content,
                    "source_file": file_path.name
                }
                print(f"  ✓ Extracted {len(raw_content)} characters")
            else:
                print(f"  ✗ Failed to process content")
        else:
            print(f"  ✗ Failed to extract content")
    
    # Generate Python code for integration
    if module_data:
        print(f"\nGenerating {output_file}...")
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("# Extracted module content from .docx files\n")
            f.write("# Generated automatically by extract_modules.py\n\n")
            f.write("MODULE_CONTENT = {\n")
            
            for module_id, data in sorted(module_data.items()):
                f.write(f'    "{module_id}": {{\n')
                f.write(f'        "title": "{data["title"]}",\n')
                f.write(f'        "source_file": "{data["source_file"]}",\n')
                f.write(f'        "content": """{data["content"]}"""\n')
                f.write(f'    }},\n')
            
            f.write("}\n")
        
        print(f"✓ Generated {output_file} with {len(module_data)} modules")
        print("\nNext steps:")
        print("1. Review the extracted content in extracted_module_content.py")
        print("2. Update user/views.py to use the extracted content")
        print("3. Test the integration in the web application")
        
    else:
        print("No content was successfully extracted from any files")

if __name__ == "__main__":
    main()
